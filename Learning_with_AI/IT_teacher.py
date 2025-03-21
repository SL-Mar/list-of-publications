import os
import openai
from docx import Document
from dotenv import load_dotenv, find_dotenv
from langchain.prompts import (
    ChatPromptTemplate,
    MessagesPlaceholder,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate
)
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from gtts import gTTS
import pygame
import speech_recognition as sr
import tempfile
import logging

# Load environment variables from .env file
load_dotenv(find_dotenv())

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set the OpenAI API key from the environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')

class ChatParticipant:
    """
    Base class for chat participants.
    """
    def __init__(self, role_syncrasy, model_name="gpt-3.5-turbo", temperature=0.5):
        self.role_syncrasy = role_syncrasy
        self.model_name = model_name
        self.temperature = temperature
        self.memory = ConversationBufferMemory(return_messages=True)
        self.chatbot = None
        self.init_chatbot()

    def init_chatbot(self):
        """
        Initializes the chatbot with the given role syncrasy and model.
        """
        prompt_template = ChatPromptTemplate.from_messages([
            SystemMessagePromptTemplate.from_template(self.role_syncrasy),
            MessagesPlaceholder(variable_name="history"),
            HumanMessagePromptTemplate.from_template("{input}")
        ])
        self.chatbot = ConversationChain(
            memory=self.memory, 
            prompt=prompt_template, 
            llm=ChatOpenAI(model_name=self.model_name, temperature=self.temperature), 
            verbose=False
        )

    def get_response(self, input_text):
        """
        Gets the response from the chatbot based on user input.
        """
        try:
            response = self.chatbot.predict(input=input_text)
            return response
        except Exception as e:
            logger.error(f"Error getting response from OpenAI: {e}")
            return "Mi dispiace, al momento non posso rispondere. Riprova più tardi."

class ItalianTeacher(ChatParticipant):
    """
    Italian language teacher chatbot.
    """
    pass  # Additional functionalities can be added here if needed

def save_conversation_to_docx(conversation_history, filename="conversation.docx"):
    """
    Saves the conversation history to a DOCX file.
    """
    try:
        doc = Document()
        doc.add_heading('Conversation History', 0)
        for role, line in conversation_history:
            p = doc.add_paragraph()
            p.add_run(f"{role}: ").bold = True
            p.add_run(line)
        doc.save(filename)
        print(f"\n📄 Conversation saved as '{filename}'.\n")
    except Exception as e:
        logger.error(f"Error saving conversation to DOCX: {e}")
        print("\n⚠️ Errore nel salvataggio della conversazione.\n")

def convert_text_to_speech(text, language='it'):
    """
    Converts text to speech using gTTS and returns the path to the audio file.
    """
    try:
        tts = gTTS(text=text, lang=language, slow=False)
        temp_audio_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        tts.save(temp_audio_file.name)
        return temp_audio_file.name
    except Exception as e:
        logger.error(f"Error in TTS conversion: {e}")
        return None

def play_audio(file_path):
    """
    Plays the audio file using pygame.
    """
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(file_path)
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        logger.error(f"Error playing audio with pygame: {e}")
        print("⚠️ Errore durante la riproduzione dell'audio.\n")

def transcribe_speech(language='it'):
    """
    Captures audio from the microphone and transcribes it to text using SpeechRecognition.
    """
    recognizer = sr.Recognizer()
    microphone = sr.Microphone()

    with microphone as source:
        print("🎤 In ascolto... Parla ora.")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)

    try:
        print("📡 Trascrizione in corso...")
        text = recognizer.recognize_google(audio, language=language)
        print(f"🧑‍🏫 Hai detto: {text}")
        return text
    except sr.UnknownValueError:
        print("❌ Non ho capito cosa hai detto. Per favore riprova.")
        return None
    except sr.RequestError as e:
        logger.error(f"Errore nel servizio di riconoscimento vocale: {e}")
        print("⚠️ Errore nel servizio di riconoscimento vocale. Riprova più tardi.")
        return None

def conduct_conversation(teacher):
    """
    Manages the conversation loop between the user and the Italian teacher.
    """
    conversation_history = []  # Initialize conversation history

    print("🎓 **Virtual Italian Teacher is now active.** 🎓")
    print("Parla direttamente al microfono per interagire.")
    print("Comandi disponibili:")
    print(" - 'END': Terminare la conversazione.")
    print(" - 'SAVE': Salvare la cronologia della conversazione.\n")

    while True:
        user_input = transcribe_speech(language='it')
        
        if user_input is None:
            continue  # Skip to next iteration if transcription failed

        if user_input.upper() == 'END':
            print("\n🛑 Terminando la conversazione. Arrivederci!\n")
            break

        if user_input.upper() == 'SAVE':
            save_conversation_to_docx(conversation_history)
            continue

        # Add user input to conversation history
        conversation_history.append(("User", user_input))

        # Get response from the teacher based on user input
        teacher_answer = teacher.get_response(user_input)
        print(f"👩‍🏫 Teacher: {teacher_answer}\n")

        # Add teacher's response to conversation history
        conversation_history.append(("Teacher", teacher_answer))

        # Convert teacher's response to speech using gTTS
        audio_file = convert_text_to_speech(teacher_answer, language='it')
        if audio_file:
            try:
                play_audio(audio_file)
            except Exception as e:
                logger.error(f"Error playing audio: {e}")
            finally:
                try:
                    os.remove(audio_file)
                except Exception as e:
                    logger.error(f"Error deleting audio file: {e}")
        else:
            print("⚠️ Failed to generate audio for the response.\n")

def main():
    """
    Main function to initialize the chatbot and start the conversation.
    """
    # Define role-specific syncrasy text for the Italian teacher
    teacher_syncrasy = '''
    Sei un paziente e competente insegnante di lingua italiana. 
    Il tuo compito è aiutare l'utente a imparare l'italiano attraverso conversazioni interattive. 
    Fornisci spiegazioni chiare, correggi i loro errori e offri esempi per illustrare i concetti linguistici. 
    Incoraggia l'utente a praticare la conversazione, l'ascolto, la lettura e la scrittura in italiano. 
    Rispondi esclusivamente in lingua italiana, a meno che non ti venga richiesto diversamente.
    '''

    # Initialize the Italian Teacher chatbot
    teacher = ItalianTeacher(teacher_syncrasy)

    # Start the conversation with the teacher
    conduct_conversation(teacher)
 
if __name__ == '__main__':
    main()
